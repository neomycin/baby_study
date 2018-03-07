import docx
import random

def math_add_sub():
    number_a = random.randint(10, 199)
    number_b = random.randint(10, 199)
    number_c = random.randint(10, 199)
    result_b = number_a + number_b - number_c
    while result_b < 0:
        number_a = random.randint(10, 199)
        number_b = random.randint(10, 199)
        number_c = random.randint(10, 199)
        result_b = number_a + number_b - number_c
    return '☐ 123 + 456 = （　）'

'''
document.add_heading(u'二级标题', 2)

# 添加文本
paragraph = document.add_paragraph(u'我们在做文本测试！')
paragraph = document.add_paragraph(u'')

# 设置字号
run = paragraph.add_run(u'设置字号、')
run.font.size = Pt(24)

# 设置字体
run = paragraph.add_run('Set Font,')
run.font.name = 'Consolas'

# 设置中文字体
run = paragraph.add_run(u'设置中文字体、')
run.font.name = u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

# 设置粗体
run = paragraph.add_run(u'粗体').bold = True

# 增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

# 增加无序列表
document.add_paragraph(
    u'无序列表元素1', style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2', style='List Bullet'
)
# 增加有序列表
document.add_paragraph(
    u'有序列表元素1', style='List Number'
)
document.add_paragraph(
    u'有序列表元素2', style='List Number'
)
# 增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
# document.add_picture('image.bmp', width=Inches(1.25))
'''