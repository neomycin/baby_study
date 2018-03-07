from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import random
import datetime

# 打开文档
# ＋－×÷  □□★……≠▲◆●☆
# math_add math_sub math_mul math_div

def math_add_sub(num):
    number_a = random.randint(10, num)
    number_b = random.randint(10, num)
    result_b = number_a - number_b
    if result_b < 0:
        math_sub_str = str(number_b) + ' － ' + str(number_a)
    else:
        math_sub_str = str(number_a) + ' － ' + str(number_b)
    math_add_str = str(number_a) + ' ＋ ' + str(number_b)
    if random.randint(0, 2) == 0:
        return '□　' + math_add_str + ' =   '
    return '□　' + math_sub_str + ' =   '

def math_add_sub_grade_two(num):
    if random.randint(0, 9) < 6:
        result_a = random.randint(10, 40) * 10
    else:
        result_a = random.randint(100, 399)
    number_a = random.randint(10, result_a)
    number_b = result_a - number_a
    number_c = random.randint(10, num)

    if result_a > number_c:
        math_str = str(number_a) + '－' + str(number_c) + '＋' + str(number_b)
    else:
        math_str = str(number_a) + '＋' + str(number_c) + '＋' + str(number_b)

    return '□　' + math_str + ' =   '

def math_add_sub_grade_three(num, d):
    number_a = random.randint(10, num)
    number_b = random.randint(10, num)
    result_a = number_a + number_b
    if random.randint(0, 9) < (9 - d):
        while result_a % 10 != 0:
            number_a = random.randint(10, num)
            number_b = random.randint(10, num)
            result_a = number_a + number_b
    if result_a > 200:
        number_c = random.randint(10, 99)
    else:
        number_c = random.randint(10, num)
    math_str = str(number_a) + '＋' + str(number_c) + '＋' + str(number_b)
    if number_c < result_a:
        if random.randint(0, 9) < 2:
            math_str = str(number_a) + '－' + str(number_c) + '＋' + str(number_b)
    return '□　' + math_str + ' =   '

def math_mul():
    number_a = random.randint(2, 9)
    number_b = random.randint(2, 9)
    result_a = number_a * number_b
    while result_a <= 10:
        number_a = random.randint(2, 9)
        number_b = random.randint(2, 9)
        result_a = number_a * number_b
    math_mul_str = str(number_a) + ' × ' + str(number_b)
    return '□　' + math_mul_str + ' =   '

def math_div():
    number_a = random.randint(2, 9)
    number_b = random.randint(2, 9)
    result_a = number_a * number_b
    while result_a <= 10:
        number_a = random.randint(2, 9)
        number_b = random.randint(2, 9)
        result_a = number_a * number_b
    math_div_str = str(result_a) + ' ÷ ' + str(number_a)
    return '□　' + math_div_str + ' =   '

def math_str():
    m = random.randint(0, 18)
    if m < 3:
        return math_add_sub(199)
    if m < 5:
        return math_mul()
    if m < 7:
        return math_div()
    return math_add_sub_grade_two(199)

def create_doc(page):
    # 打开文档
    document = Document()
    j = 0
    while j < page:
        # 加入不同等级的标题
        document.add_heading(u'计算小超市　　　　　　　　月　日 ', 0)
        # 增加表格

        document.add_heading(u'基础过关题：', 1)
        document.add_heading(u' ', 2)
        table = document.add_table(rows=3, cols=3)
        i = 0
        while i < 3:
            hdr_cells = table.rows[i].cells
            hdr_cells[0].text = (math_add_sub(99))
            hdr_cells[1].text = (math_add_sub(99))
            hdr_cells[2].text = (math_add_sub(99))
            i = i + 1

        document.add_heading(u'能力提高题：', 1)
        document.add_heading(u' ', 2)
        table = document.add_table(rows=12, cols=3)
        i = 0
        while i < 12:
            hdr_cells = table.rows[i].cells
            hdr_cells[0].text = (math_str())
            hdr_cells[1].text = (math_str())
            hdr_cells[2].text = (math_str())
            i = i + 1
        # 增加分页
        if j < page - 1:
            document.add_page_break()
        j = j + 1
    # 保存文件
    now = datetime.datetime.now()

    filename = now.strftime('%Y%m%d%H%M%S')
    document.save('D:\\数学' + filename + '.docx')

print(math_add_sub_grade_two(199))
print(math_add_sub_grade_two(199))

print(math_add_sub_grade_three(199, 6))
print(math_mul())
print(math_div())

#生成多少页计算文档
create_doc(7)
