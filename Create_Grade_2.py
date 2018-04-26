from docx import Document
import random
import datetime

# 打开文档
# ＋－×÷  □□★……≠▲◆●☆
# math_add math_sub math_mul math_div

def math_split(num):
    math_split_str = ''
    number = random.randint(4, 9) * 10
    #print(str(number))
    i = num
    while i > 1:
        number_random = random.randint(1, 4)
        #print(str(number_random))
        if random.randint(0, 1) == 0:
            math_split_str = math_split_str + str(number + number_random) + ' + '
        else:
            math_split_str = math_split_str + str(number - number_random) + ' + '
        i = i - 1
    math_split_str = math_split_str + str(number - random.randint(1, 4))
    return math_split_str + ' = '

def create_doc(page):
    # 打开文档
    document = Document()
    j = 0
    while j < page:
        # 加入不同等级的标题
        document.add_heading(u'计算小超市　　　　　　　　月　日 ', 0)
        # 增加表格

        document.add_heading(u'基础过关题：', 1)
        document.add_heading(u' ', 2)
        table = document.add_table(rows=6, cols=2)
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = math_split(6)


        document.add_page_break()
        j = j + 1
    # 保存文件
    now = datetime.datetime.now()

    filename = now.strftime('%Y%m%d%H%M%S')
    document.save('D:\\数学' + filename + '.docx')

#生成多少页计算文档
create_doc(1)
